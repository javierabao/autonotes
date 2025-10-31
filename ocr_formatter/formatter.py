from docx import Document
from transformers import AutoModelForCausalLM, AutoTokenizer
import re

class ModelManager:
    _instance = None
    
    @classmethod
    def get_instance(cls):
        if cls._instance is None:
            cls._instance = cls._load_model()
        return cls._instance
    
    @staticmethod
    def _load_model():
        # Using Qwen model for efficient text formatting
        model_name = "Qwen/Qwen3-0.6B"
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForCausalLM.from_pretrained(
            model_name,
            torch_dtype="auto",
            device_map="auto"
        )
        return {"model": model, "tokenizer": tokenizer}

class TextFormatter:
    def __init__(self):
        """Initialize the text formatter with T5 model."""
        self.model_manager = ModelManager.get_instance()
        
    def format_with_llm(self, text: str) -> str:
        """
        Format text using Qwen model.
        
        Args:
            text: Input text to format
            
        Returns:
            str: Formatted text
        """
        # Create the system and user messages
        messages = [
            {
                "role": "system",
                "content": """You are an assistant that formats text extracted from class lectures into markdown notes. Structure the text sections properly with headers, bullet points, and code blocks where appropriate.
Restrictions:
1. If there are missing words fill the gaps, but do not change the grammar nor words unless necessary to fix OCR errors.
2. If there are invented words, correct them appropriately.
3. If there are non-sense characters, remove them.
4. Do not add any additional commentary or remarks, just provide the formatted text.
Here is the text:"""
            },
            {
                "role": "user",
                "content": text
            }
        ]
        
        # Get model and tokenizer
        model = self.model_manager["model"]
        tokenizer = self.model_manager["tokenizer"]
        
        # Convert messages to chat template format
        input_ids = tokenizer.apply_chat_template(
            messages,
            return_tensors="pt",
            max_length=2048,
            truncation=True
        ).to(model.device)
        
        # Generate response
        outputs = model.generate(
            input_ids,
            max_length=2048,
            num_beams=4,
            temperature=0.01,
            do_sample=True,
            eos_token_id=tokenizer.eos_token_id,
            pad_token_id=tokenizer.pad_token_id
        )
        
        output_ids = outputs[0][len(input_ids[0]):].tolist() 
        formatted_text = tokenizer.decode(output_ids, skip_special_tokens=True)
        # Remove any <think> artifacts produced by the LLM.
        # If both opening and closing tags exist, remove the tags and everything between them.
        # Otherwise remove any lone tag left behind.
        try:
            if re.search(r"<think>.*?</think>", formatted_text, flags=re.I | re.S):
                formatted_text = re.sub(r"<think>.*?</think>", "", formatted_text, flags=re.I | re.S)
            else:
                # remove lone opening or closing tags if present
                formatted_text = re.sub(r"</?think>", "", formatted_text, flags=re.I)
        except Exception:
            # Be conservative: on regex failures, fallback to simple replace
            formatted_text = formatted_text.replace("<think>", "").replace("</think>", "")

        # Trim leftover whitespace/newlines caused by removals
        formatted_text = formatted_text.strip()

        return formatted_text

    def save_as_markdown(self, text: str, output_path: str):
        """
        Save formatted text as a Markdown file.
        
        Args:
            text: Formatted text
            output_path: Path to save the markdown file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    def save_as_docx(self, text: str, output_path: str):
        """
        Save formatted text as a DOCX file, preserving basic markdown formatting.
        
        Args:
            text: Formatted text (in markdown format)
            output_path: Path to save the DOCX file
        """
        # Create a new Document
        doc = Document()
        
        # Split text into paragraphs
        paragraphs = text.split('\n')
        
        for line in paragraphs:
            line = line.rstrip()
            
            # Skip empty lines
            if not line:
                doc.add_paragraph()
                continue
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                doc.add_heading(title, level=min(level, 9))
                continue
            
            # Handle list items
            if line.strip().startswith(('- ', '* ')):
                p = doc.add_paragraph(style='List Bullet')
                text = line.lstrip('- *').strip()
                # Handle bold text in list items
                if '**' in text:
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            run.bold = True
                else:
                    p.add_run(text)
                continue
            
            # Handle regular paragraphs with potential bold text
            p = doc.add_paragraph()
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
            else:
                p.add_run(line)
        
        # Save the document with error handling
        try:
            doc.save(output_path)
        except Exception as e:
            import logging
            logging.error(f"Error saving DOCX file: {str(e)}")
            raise